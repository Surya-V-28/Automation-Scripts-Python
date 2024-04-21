import re
from docx import Document
from flask import Flask, render_template, request, send_file
from io import BytesIO

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

#Find the LinkedIn Id of the Person 
def find_first_linkedin(text):
    # Regular expression to match LinkedIn profile URLs
    linkedin_regex = r"(?:(?:http|https):\/\/)?(?:www\.)?linkedin\.com\/(?:in|profile)\/[\w-]+"
    match = re.search(linkedin_regex, text)
    if match:
        return match.group()
    return None

#Find the Email if for the Person
def find_first_email_id(text):
    email_regex =  r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    match = re.search(email_regex,text)
    if(match):
        return match.group()
    return None

#Finding the Phone for the person
def find_first_phone_number(text):
    # Regular expression to match phone numbers in various formats
    phone_regex = r"\b(?:\+\d{1,2}\s*)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"
    match = re.search(phone_regex, text)
    if match:
        print(match,"found happening")
        return match.group()
    return None


#Copying the  Original Styles of the Person
def copy_paragraph_style(source_paragraph, target_paragraph):
    # Copy the style from the source paragraph to the target paragraph
    target_paragraph.style = source_paragraph.style

#Date  and Month Full form Replacement Works
def replace_words(paragraph):
    # Replace words in the paragraph
    replacements = {
            "Jan": "January",
            "Feb": "February",
            "Mar": "March",
            "Apr": "April",
            "May": "May",
            "Jun": "June",
            "Jul": "July",
            "Aug": "August",
            "Sep": "September",
            "Oct": "October",
            "Nov": "November",
            "Dec": "December"
    }
    for original, replacement in replacements.items():
        if original in paragraph.text:
            if replacement not in paragraph.text:
                paragraph.text = paragraph.text.replace(original, replacement)


def convert_date_ranges(text):
    # Regular expression to match date ranges like '98-02' or '22-24'
    date_range_regex = date_range_regex = r"(\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{2,4})\s*-\s*(\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{2,4})"
    converted_text = re.sub(date_range_regex, r"20\1 - 20\2", text)
    return converted_text


#Main Block of the Code
def process_document(docx_file):
    # Load the original document
    original_doc = Document(docx_file)
    
    # Create a new document
    new_doc = Document()

    # Find the first phone number in the document
    first_phone_number = None
    first_email = None
    first_linked_in = None
    for paragraph in original_doc.paragraphs:
        phone_number = find_first_phone_number(paragraph.text)
        if phone_number:
            first_phone_number = phone_number
            break
   
   #Adding the Email to the new Document
    for paragraph in original_doc.paragraphs:
        email = find_first_email_id(paragraph.text) 
        if email:
            first_email = email
            break
    

    # Adding the Linked Id to the New Document if needed
    for paragraph in original_doc.paragraphs:
        linkedIn = find_first_linkedin(paragraph.text)
        if linkedIn: 
            first_linked_in = linkedIn
            break
        

    # Add the first phone number to the start of the new document
    if first_phone_number:
        new_doc.add_paragraph(first_phone_number)
   
    # Add a separator
    new_doc.add_paragraph("--------------------")
    if first_email:
        new_doc.add_paragraph(first_email)
    # Add a separator
    new_doc.add_paragraph("--------------------")

    if first_linked_in:
        new_doc.add_paragraph(first_linked_in)
    # Add a separator
    new_doc.add_paragraph("--------------------")
    contact_info = f"Phone: {phone_number} | Email: {first_email} | LinkedIn: {first_linked_in}"
    new_doc.add_paragraph(contact_info)


    # Convert date ranges like '98-02' or '22-24' to '1998 - 2002' or '2022 - 2024'
    for paragraph in original_doc.paragraphs:
        converted_text = convert_date_ranges(paragraph.text)
        paragraph.text = converted_text
    
    # Copy paragraphs and replace words
    for paragraph in original_doc.paragraphs:
        new_paragraph = new_doc.add_paragraph('')
        copy_paragraph_style(paragraph, new_paragraph)
        new_paragraph.text = paragraph.text
        replace_words(new_paragraph)

    
    # Save the new document to a BytesIO object
    output = BytesIO()
    new_doc.save(output)
    output.seek(0)
    return output

@app.route('/process_document', methods=['POST'])
def handle_process_document():
    if 'file' not in request.files:
        return "No file part"
    
    file = request.files['file']
    
    if file.filename == '':
        return "No selected file"
    
    output = process_document(file)
    return send_file(output, as_attachment=True, download_name="output.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True)
